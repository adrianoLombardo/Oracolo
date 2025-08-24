from __future__ import annotations

import argparse
import importlib
import json
import logging
import shutil
import sys
import warnings
from pathlib import Path
from typing import Iterable, List, Optional

# --- Assicura che la root del progetto sia in sys.path anche se lanci come file ---
ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# Dipendenze opzionali per l'estrazione testo
try:
    from pypdf import PdfReader  # pip install pypdf
except Exception:
    PdfReader = None

try:
    import docx  # pip install python-docx
except Exception:
    docx = None

# language detection opzionale
try:
    from langdetect import detect
except Exception:
    detect = None

# settings opzionali
try:
    from src.config import Settings  # se disponibile
except Exception:
    Settings = None  # fallback

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")


# ----------------------------- Simple fallback DB ----------------------------- #
class SimpleDocumentDB:
    """
    Docstore minimale basato su JSON.

    Struttura attesa: {"documents": [{"id": "...", "text": "...", "topic": "..."}]}
    Nota: salva il testo nell'indice (ok per volumi piccoli/medi).
    """

    def __init__(self, index_path: str | Path) -> None:
        self.index_path = Path(index_path)
        self.index_path.parent.mkdir(parents=True, exist_ok=True)
        self._data = {"documents": []}
        self._load()

    def _load(self) -> None:
        if self.index_path.exists():
            try:
                self._data = json.loads(self.index_path.read_text(encoding="utf-8"))
                if "documents" in self._data and isinstance(self._data["documents"], list):
                    return
                # compatibilità con vecchi indici {"docs": []}
                if "docs" in self._data and isinstance(self._data["docs"], list):
                    self._data = {"documents": self._data["docs"]}
                else:
                    self._data = {"documents": []}
            except Exception:
                self._data = {"documents": []}

    def _save(self) -> None:
        self.index_path.write_text(
            json.dumps(self._data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def add_documents(self, documents: List[dict]) -> None:
        # sostituisce per id
        existing = {d["id"]: i for i, d in enumerate(self._data["documents"]) if "id" in d}
        for doc in documents:
            doc_id = doc.get("id")
            if not doc_id:
                continue
            if doc_id in existing:
                self._data["documents"][existing[doc_id]] = doc
            else:
                self._data["documents"].append(doc)
        self._save()

    def delete_documents(self, ids: List[str]) -> None:
        self._data["documents"] = [d for d in self._data["documents"] if d.get("id") not in ids]
        self._save()

    # Comodo per il bottone "Aggiorna indice": rilegge i file dai path id
    def reindex(self) -> None:
        new_docs = []
        for d in self._data["documents"]:
            fid = d.get("id")
            if not fid:
                continue
            p = Path(fid)
            if p.exists() and p.is_file():
                txt = _read_file_text(p)
                new_docs.append({"id": str(p), "text": txt, "topic": d.get("topic")})
        self._data["documents"] = new_docs
        self._save()


# ----------------------------- Utilità I/O ----------------------------------- #
ALLOWED_EXTS = {".pdf", ".docx", ".txt", ".md"}


def _gather_files(paths: Iterable[str]) -> List[Path]:
    files: List[Path] = []
    for raw in paths:
        path = Path(raw)
        if path.is_dir():
            for file in path.rglob("*"):
                if file.is_file() and file.suffix.lower() in ALLOWED_EXTS:
                    files.append(file)
        elif path.is_file():
            if path.suffix.lower() in ALLOWED_EXTS:
                files.append(path)
            else:
                logging.warning("Skipping unsupported file %s", path.name)
        else:
            logging.warning("Skipping unknown path %s", path)
    # de-dup per percorso assoluto
    uniq: List[Path] = []
    seen = set()
    for f in files:
        s = str(f.resolve())
        if s not in seen:
            uniq.append(f)
            seen.add(s)
    return uniq


def _extract_pdf_text(pdf_path: Path) -> str:
    if PdfReader is None:
        logging.warning("pypdf non installato: salto estrazione testo per %s", pdf_path.name)
        return ""
    try:
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", category=UserWarning)
            reader = PdfReader(str(pdf_path), strict=False)
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n".join(parts).strip()
        if not text:
            logging.info("Nessun testo estratto da %s", pdf_path.name)
        return text
    except Exception as e:
        logging.error("PDF read error for %s: %s", pdf_path.name, e)
        return ""


def _extract_docx_text(docx_path: Path) -> str:
    if docx is None:
        logging.warning("python-docx non installato: salto estrazione testo per %s", docx_path.name)
        return ""
    try:
        d = docx.Document(str(docx_path))
        return "\n".join(p.text for p in d.paragraphs).strip()
    except Exception as e:
        logging.error("DOCX read error for %s: %s", docx_path.name, e)
        return ""


def _read_file_text(file: Path) -> str:
    ext = file.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf_text(file)
    if ext == ".docx":
        return _extract_docx_text(file)
    if ext in {".txt", ".md"}:
        try:
            return file.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            logging.error("Text read error for %s: %s", file.name, e)
            return ""
    return ""


def _extract_title(file: Path) -> str:
    ext = file.suffix.lower()
    if ext == ".pdf" and PdfReader is not None:
        try:
            with warnings.catch_warnings():
                warnings.filterwarnings("ignore", category=UserWarning)
                reader = PdfReader(str(file), strict=False)
            meta = getattr(reader, "metadata", None)
            if meta:
                title = getattr(meta, "title", "") or meta.get("/Title", "")
                if isinstance(title, str) and title.strip():
                    return title.strip()
        except Exception:
            pass
    if ext == ".docx" and docx is not None:
        try:
            d = docx.Document(str(file))
            title = getattr(d.core_properties, "title", "")
            if title:
                return title
        except Exception:
            pass
    return file.stem


def _detect_language(text: str) -> str:
    if detect is None:
        return ""
    try:
        return detect(text)
    except Exception:
        return ""


# ----------------------------- DB loader ------------------------------------- #
def _load_db(path: str) -> object:
    # prova a caricare un DocumentDB custom
    try:
        module = importlib.import_module("documentdb")
        DocumentDB = getattr(module, "DocumentDB")
        logging.info("Uso DocumentDB personalizzato (%s)", module.__file__)
        return DocumentDB(path)
    except Exception:
        logging.info("Uso SimpleDocumentDB (fallback) → %s", path)
        return SimpleDocumentDB(path)


# ----------------------------- Operazioni ------------------------------------ #
def _add(paths: Iterable[str], docstore_path: str, topic: str | None = None) -> None:
    db = _load_db(docstore_path)
    files = _gather_files(paths)
    documents = []
    for file in files:
        text = _read_file_text(file)
        title = _extract_title(file)
        lang = _detect_language(text)
        doc = {"id": str(file), "text": text, "title": title, "lang": lang}
        if topic:
            doc["topic"] = topic
        documents.append(doc)
    if documents:
        db.add_documents(documents)
    logging.info("Indexed %d document(s)", len(documents))


def _remove(paths: Iterable[str], docstore_path: str) -> None:
    db = _load_db(docstore_path)
    files = _gather_files(paths)
    ids = [str(file) for file in files]
    # compat con nomi alternativi
    if hasattr(db, "delete_documents"):
        db.delete_documents(ids)
    elif hasattr(db, "remove_documents"):
        db.remove_documents(ids)
    else:
        raise AttributeError("DocumentDB does not support removing documents")
    logging.info("Removed %d document(s)", len(ids))


def _reindex(docstore_path: str) -> None:
    db = _load_db(docstore_path)
    if hasattr(db, "reindex"):
        db.reindex()
        logging.info("Reindex completed.")
    else:
        logging.warning("DocumentDB non supporta reindex().")


def _clear(docstore_path: str) -> None:
    """Svuota completamente il DocumentDB."""
    db = _load_db(docstore_path)
    if hasattr(db, "clear"):
        db.clear()  # type: ignore[call-arg]
    elif isinstance(db, SimpleDocumentDB):
        db._data = {"documents": []}
        db._save()
    else:
        index_path = Path(docstore_path)
        index_path.write_text(
            json.dumps({"documents": []}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    logging.info("Cleared document store.")


# ----------------------------- Entry point ------------------------------------ #
def _default_docstore_path() -> str:
    # Se hai Settings, prova a usarlo; altrimenti default sensato
    if Settings is not None:
        try:
            setts = Settings.model_validate_yaml(ROOT / "settings.yaml")
            path = getattr(setts, "docstore_path", None)
            if path:
                return str(path)
        except Exception:
            pass
    # fallback
    return str(ROOT / "data" / "index.json")


def main() -> None:
    parser = argparse.ArgumentParser(description="Ingest or remove documents")
    parser.add_argument(
        "--path",
        default=_default_docstore_path(),
        help="Path al document store (default: settings.yaml docstore_path o data/index.json)",
    )
    parser.add_argument("--topic", help="Etichetta topic da associare ai documenti aggiunti")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--add", nargs="+", help="File o cartelle da indicizzare (PDF/DOCX/TXT/MD)")
    group.add_argument("--remove", nargs="+", help="File o cartelle da rimuovere dall'indice")
    group.add_argument("--reindex", action="store_true", help="Rigenera l'indice rileggendo i file già noti")
    group.add_argument("--clear", action="store_true", help="Svuota completamente l’indice")
    args = parser.parse_args()

    if args.add:
        _add(args.add, args.path, args.topic)
    elif args.remove:
        _remove(args.remove, args.path)
    elif args.reindex:
        _reindex(args.path)
    elif args.clear:
        _clear(args.path)


if __name__ == "__main__":
    main()
